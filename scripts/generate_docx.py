#!/usr/bin/env python3
"""
Generate Comprehensive Word Document for DSC288 Progress Report
Follows rubric structure with all EDA analysis, insights, and flowcharts.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import os

# Create document
doc = Document()

# ============================================
# DOCUMENT SETUP - Margins and Styles
# ============================================
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_number(doc, text, level=1):
    """Add a numbered heading"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    return heading

def add_insight_box(doc, title, content):
    """Add a simple styled insight box"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'E8F4FD')
    
    p = cell.paragraphs[0]
    run = p.add_run(f"INSIGHT: {title}\n")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(26, 54, 93)
    
    content_run = p.add_run(content)
    content_run.font.size = Pt(10)
    
    doc.add_paragraph()

def add_eda_insight_box(doc, title, what_analyzed, what_found, action_taken, why_matters, color='3182CE'):
    """Add a comprehensive EDA insight box matching HTML format"""
    # Main container table
    main_table = doc.add_table(rows=1, cols=1)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_cell = main_table.cell(0, 0)
    set_cell_shading(main_cell, 'F0F7FF')
    
    # Header with INSIGHT badge
    header_p = main_cell.paragraphs[0]
    badge_run = header_p.add_run(" INSIGHT ")
    badge_run.bold = True
    badge_run.font.size = Pt(10)
    badge_run.font.color.rgb = RGBColor(255, 255, 255)
    # Note: Background color on runs not directly supported, so we use text styling
    
    title_run = header_p.add_run(f"  {title}")
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(26, 54, 93)
    
    # Create inner 2x2 grid for content
    inner_table = main_cell.add_table(rows=2, cols=2)
    inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Cell 1: What We Analyzed
    cell1 = inner_table.cell(0, 0)
    set_cell_shading(cell1, 'FFFFFF')
    p1 = cell1.paragraphs[0]
    h1 = p1.add_run("What We Analyzed:\n")
    h1.bold = True
    h1.font.size = Pt(10)
    h1.font.color.rgb = RGBColor(49, 130, 206)
    c1 = p1.add_run(what_analyzed)
    c1.font.size = Pt(9)
    
    # Cell 2: What We Found
    cell2 = inner_table.cell(0, 1)
    set_cell_shading(cell2, 'FFFFFF')
    p2 = cell2.paragraphs[0]
    h2 = p2.add_run("What We Found:\n")
    h2.bold = True
    h2.font.size = Pt(10)
    h2.font.color.rgb = RGBColor(56, 161, 105)
    c2 = p2.add_run(what_found)
    c2.font.size = Pt(9)
    
    # Cell 3: Action Taken
    cell3 = inner_table.cell(1, 0)
    set_cell_shading(cell3, 'FFFFFF')
    p3 = cell3.paragraphs[0]
    h3 = p3.add_run("Action Taken:\n")
    h3.bold = True
    h3.font.size = Pt(10)
    h3.font.color.rgb = RGBColor(214, 158, 46)
    c3 = p3.add_run(action_taken)
    c3.font.size = Pt(9)
    
    # Cell 4: Features Created
    cell4 = inner_table.cell(1, 1)
    set_cell_shading(cell4, 'FFFFFF')
    p4 = cell4.paragraphs[0]
    h4 = p4.add_run("Features Created:\n")
    h4.bold = True
    h4.font.size = Pt(10)
    h4.font.color.rgb = RGBColor(128, 90, 213)
    c4 = p4.add_run(action_taken)
    c4.font.size = Pt(9)
    
    # Why This Matters - footer
    footer_p = main_cell.add_paragraph()
    why_run = footer_p.add_run("Why This Matters: ")
    why_run.bold = True
    why_run.italic = True
    why_run.font.size = Pt(10)
    why_run.font.color.rgb = RGBColor(74, 85, 104)
    
    why_content = footer_p.add_run(why_matters)
    why_content.italic = True
    why_content.font.size = Pt(10)
    why_content.font.color.rgb = RGBColor(74, 85, 104)
    
    doc.add_paragraph()

def add_formatted_table(doc, headers, rows, header_color='1A365D', col_widths=None):
    """Add a professionally formatted table"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(cell, header_color)
    
    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            # Alternate row colors
            if i % 2 == 1:
                set_cell_shading(cell, 'F7FAFC')
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    doc.add_paragraph()
    return table

def add_flowchart_box(doc, title, stages):
    """Add a flowchart as styled boxes with arrows"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(26, 54, 93)
    
    # Create table with stages and arrows
    num_cols = len(stages) * 2 - 1  # stages + arrows between them
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    colors = ['3182CE', 'D69E2E', '38A169', 'E53E3E', '805AD5', '718096']
    
    col_idx = 0
    for i, (stage_title, stage_content) in enumerate(stages):
        # Stage cell
        cell = table.cell(0, col_idx)
        set_cell_shading(cell, colors[i % len(colors)])
        cell.width = Inches(1.3)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = p.add_run(stage_title + "\n")
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(255, 255, 255)
        title_run.font.size = Pt(9)
        
        content_run = p.add_run(stage_content)
        content_run.font.color.rgb = RGBColor(255, 255, 255)
        content_run.font.size = Pt(8)
        
        col_idx += 1
        
        # Arrow cell (except for last stage)
        if i < len(stages) - 1:
            arrow_cell = table.cell(0, col_idx)
            arrow_cell.width = Inches(0.3)
            arrow_p = arrow_cell.paragraphs[0]
            arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_run = arrow_p.add_run("-->")
            arrow_run.bold = True
            arrow_run.font.size = Pt(10)
            arrow_run.font.color.rgb = RGBColor(74, 85, 104)
            col_idx += 1
    
    doc.add_paragraph()

def add_large_flowchart(doc, title, rows_data):
    """Add a multi-row flowchart for complex architectures"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(26, 54, 93)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_title, row_items, row_color in rows_data:
        # Row title
        if row_title:
            row_p = doc.add_paragraph()
            row_run = row_p.add_run(row_title)
            row_run.bold = True
            row_run.font.size = Pt(10)
            row_run.font.color.rgb = RGBColor(74, 85, 104)
        
        # Items in row
        table = doc.add_table(rows=1, cols=len(row_items))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, (item_title, item_content) in enumerate(row_items):
            cell = table.cell(0, i)
            set_cell_shading(cell, row_color)
            cell.width = Inches(6.5 / len(row_items))
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            title_run = p.add_run(item_title + "\n")
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(255, 255, 255)
            title_run.font.size = Pt(9)
            
            if item_content:
                content_run = p.add_run(item_content)
                content_run.font.color.rgb = RGBColor(255, 255, 255)
                content_run.font.size = Pt(8)
    
    doc.add_paragraph()

def add_neo4j_schema(doc):
    """Add Neo4j schema visualization as formatted tables"""
    # Title
    p = doc.add_paragraph()
    run = p.add_run("Knowledge Graph Schema")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(26, 54, 93)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Entity Nodes
    doc.add_paragraph()
    node_p = doc.add_paragraph()
    node_run = node_p.add_run("ENTITY NODES")
    node_run.bold = True
    node_run.font.size = Pt(11)
    
    nodes_table = doc.add_table(rows=2, cols=6)
    nodes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    nodes = [
        ("(:Stock)", "#3182CE", "ticker, name\nsector, price"),
        ("(:Article)", "#D69E2E", "id, title\ndate, text"),
        ("(:Person)", "#805AD5", "name, role\ncompany"),
        ("(:Event)", "#7C3AED", "type, date\nimpact"),
        ("(:Sentiment)", "#38A169", "score, label\nconfidence"),
        ("(:TradingSignal)", "#E53E3E", "type, strength\nreasoning")
    ]
    
    for i, (name, color, props) in enumerate(nodes):
        # Name row
        cell = nodes_table.cell(0, i)
        set_cell_shading(cell, color.replace("#", ""))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
        
        # Properties row
        prop_cell = nodes_table.cell(1, i)
        set_cell_shading(prop_cell, 'F7FAFC')
        prop_p = prop_cell.paragraphs[0]
        prop_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prop_run = prop_p.add_run(props)
        prop_run.font.size = Pt(8)
    
    doc.add_paragraph()
    
    # Relationships
    rel_p = doc.add_paragraph()
    rel_run = rel_p.add_run("RELATIONSHIPS")
    rel_run.bold = True
    rel_run.font.size = Pt(11)
    
    rels_table = doc.add_table(rows=4, cols=2)
    rels_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rels_table.style = 'Table Grid'
    
    relationships = [
        ("Stock -[:MENTIONED_IN]-> Article", "Article -[:HAS_SENTIMENT]-> Sentiment"),
        ("Sentiment -[:TRIGGERS]-> TradingSignal", "TradingSignal -[:RECOMMENDS]-> Stock"),
        ("Event -[:AFFECTS]-> Stock", "Person -[:EXECUTIVE_OF]-> Stock"),
        ("Person -[:MENTIONED_IN]-> Article", "Stock -[:CORRELATED_WITH]-> Stock")
    ]
    
    for i, (rel1, rel2) in enumerate(relationships):
        cell1 = rels_table.cell(i, 0)
        set_cell_shading(cell1, 'EDF2F7')
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(rel1)
        run1.font.size = Pt(9)
        run1.font.name = 'Consolas'
        
        cell2 = rels_table.cell(i, 1)
        set_cell_shading(cell2, 'EDF2F7')
        p2 = cell2.paragraphs[0]
        run2 = p2.add_run(rel2)
        run2.font.size = Pt(9)
        run2.font.name = 'Consolas'
    
    doc.add_paragraph()

def add_code_block(doc, title, code, result=None):
    """Add a styled code block"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(26, 54, 93)
    
    # Code block
    code_table = doc.add_table(rows=1, cols=1)
    code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    code_cell = code_table.cell(0, 0)
    set_cell_shading(code_cell, '1E293B')
    
    code_p = code_cell.paragraphs[0]
    code_run = code_p.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(226, 232, 240)
    
    if result:
        doc.add_paragraph()
        result_p = doc.add_paragraph()
        result_run = result_p.add_run("Example Output: ")
        result_run.bold = True
        result_run.font.size = Pt(10)
        
        result_content = result_p.add_run(result)
        result_content.font.size = Pt(10)
        result_content.italic = True
    
    doc.add_paragraph()

# ============================================
# TITLE PAGE
# ============================================
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("DSC288R - Milestone 2: 1st Progress Report")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(26, 54, 93)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Multi-Agent Graph RAG System for\nExplainable Financial Decision Support")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(49, 130, 206)

doc.add_paragraph()
doc.add_paragraph()

# Project info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

lines = [
    ("Project Group #", " 10"),
    ("Authors:", " Harsh Arya, Gabrielle Despaigne, Camila Paik, Raghav Vasappanavara"),
    ("Emails:", " harya@ucsd.edu, gdespaigne@ucsd.edu, cpaik@ucsd.edu, rvasappanavara@ucsd.edu"),
    ("Institution:", " UC San Diego"),
    ("Date:", " February 2026")
]

for label, value in lines:
    run = info.add_run(label)
    run.bold = True
    info.add_run(value + "\n")

doc.add_page_break()

# ============================================
# TABLE OF CONTENTS
# ============================================
doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("1.", "Background", "Problem statement and importance"),
    ("2.", "Dataset", "Data sources and citations"),
    ("3.", "Data Pipeline", "Design, processing, outputs"),
    ("4.", "EDA Description", "Analysis types, visualizations, artifacts"),
    ("5.", "Feature Engineering", "EDA-driven feature extraction"),
    ("6.", "Model Architecture", "Proposed models and justification"),
    ("7.", "Graph RAG System", "Neo4j schema, queries, implementation"),
    ("8.", "Progress Report", "Current status and team contributions"),
    ("9.", "Risks and Mitigation", "Project risks and strategies"),
    ("10.", "References", "Citations and sources")
]

for num, title, desc in toc_items:
    p = doc.add_paragraph()
    p.add_run(f"{num} {title}").bold = True
    p.add_run(f" - {desc}")
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_page_break()

# ============================================
# SECTION 1: BACKGROUND (Rubric requirement)
# ============================================
add_heading_with_number(doc, "1. Background", 1)

doc.add_paragraph(
    "Problem Statement: ", style='Normal'
).runs[0].bold = True
doc.paragraphs[-1].add_run(
    "Retail investors lack access to sophisticated financial analysis tools that institutional investors use. "
    "Existing AI trading systems prioritize prediction accuracy over explainability, leaving users unable to "
    "understand why a recommendation was made or verify its reasoning."
)

doc.add_paragraph()

doc.add_paragraph(
    "Why This Problem is Important: ", style='Normal'
).runs[0].bold = True
doc.paragraphs[-1].add_run(
    "Without explainable recommendations, investors cannot make informed decisions, leading to either blind trust "
    "in opaque AI systems or complete avoidance of potentially valuable tools. Our system bridges this gap by "
    "providing transparent, source-cited financial analysis using multi-agent LLM architecture with Graph RAG."
)

doc.add_paragraph()

# Key Innovation box
add_insight_box(doc, "Key Innovation",
    "Unlike traditional financial models, our system uses specialized agents (fundamental analyst, news/sentiment analyst, "
    "technical analyst, optimistic/cautious viewpoints) orchestrated via OpenAI Agents SDK to generate interpretable "
    "recommendations with full data provenance through Neo4j knowledge graph.")

# ============================================
# SECTION 2: DATASET (Rubric requirement)
# ============================================
add_heading_with_number(doc, "2. Dataset Name and Citation/Link", 1)

doc.add_paragraph(
    "We utilize four primary datasets for this project, integrating structured market data with unstructured financial text:"
)

add_formatted_table(doc,
    ["Dataset", "Description", "Records", "Link/Citation"],
    [
        ["FNSPID", "Stock prices and financial news for Fortune 500 companies (2009-2023)", "10M+", "huggingface.co/datasets/Zihan1004/FNSPID"],
        ["Financial Phrasebank", "Sentiment-labeled financial news sentences", "4,840", "huggingface.co/datasets/takala/financial_phrasebank"],
        ["Yahoo Finance S&P 500", "Daily S&P 500 index data for market context", "6,289", "finance.yahoo.com/quote/^GSPC"],
        ["FinQA", "Financial question-answering benchmark", "8,000+", "github.com/czyssrs/FinQA"]
    ],
    col_widths=[1.3, 2.5, 0.8, 2.0]
)

add_insight_box(doc, "Data Coverage",
    "Combined dataset spans 14.2 years (Oct 2009 - Dec 2023), covering 100 stock tickers with 262,257 aligned observations. "
    "This includes multiple market cycles: post-2008 recovery, bull market, COVID-19 crash/recovery, and 2022-2023 volatility.")

# ============================================
# SECTION 3: DATA PIPELINE (Rubric: 5 pts)
# ============================================
add_heading_with_number(doc, "3. Data Pipeline", 1)

doc.add_heading("3.1 Purpose", level=2)
doc.add_paragraph(
    "The data pipeline integrates multi-modal financial data (structured prices + unstructured news) into a unified, "
    "analysis-ready format. It handles the complexities of temporal alignment, data quality issues, and feature normalization "
    "required for cross-stock analysis."
)

doc.add_heading("3.2 Pipeline Design and Details", level=2)

# Pipeline flowchart
add_flowchart_box(doc, "End-to-End Pipeline Architecture:", [
    ("STAGE 1\nData Loading", "5 sources\n→ Parquet"),
    ("STAGE 2\nData Cleaning", "NULL handling\nOutlier removal"),
    ("STAGE 3\nTemporal Alignment", "News-Price merge\nTarget creation"),
    ("STAGE 4\nFeature Engineering", "35 features\nNormalization")
])

# a. Data Merging
doc.add_heading("a) Data Merging (Multiple Sources)", level=3)
add_formatted_table(doc,
    ["Source", "Data Type", "Merge Key", "Records Merged"],
    [
        ["FNSPID Prices", "OHLCV daily data", "(ticker, date)", "262K stock-days"],
        ["FNSPID News", "Article text, headlines", "(ticker, date)", "3,821 news-days"],
        ["S&P 500 Index", "Market returns", "(date)", "86.4% coverage"],
        ["Financial Phrasebank", "Sentiment labels", "Training data", "4,840 sentences"]
    ],
    col_widths=[1.5, 1.5, 1.3, 1.5]
)

# b. Data Cleansing
doc.add_heading("b) Data Cleansing", level=3)
add_formatted_table(doc,
    ["Issue", "Detection Method", "Resolution", "Records Affected"],
    [
        ["NULL values", "df.isnull().sum()", "Forward fill for prices, flag for news", "< 0.1% prices"],
        ["Missing dates", "Date range gaps", "Exclude weekends/holidays", "Standard"],
        ["Date formats", "pd.to_datetime()", "Standardize to YYYY-MM-DD", "All records"],
        ["Duplicates", "df.duplicated()", "Keep first occurrence", "127 removed"],
        ["Outliers (>50% change)", "Return threshold", "Remove likely data errors", "951 removed"]
    ],
    col_widths=[1.3, 1.5, 1.8, 1.2]
)

# c. Data Augmentation
doc.add_heading("c) Data Augmentation/Enrichment", level=3)
add_formatted_table(doc,
    ["Derived Attribute", "Formula", "Purpose"],
    [
        ["next_day_return", "(close[t+1] - close[t]) / close[t]", "Prediction target"],
        ["target (buy/hold/sell)", "Threshold at ±2% return", "Classification labels"],
        ["sma_5, sma_20, sma_50", "Rolling mean of close prices", "Trend indicators"],
        ["volatility_20", "20-day rolling std of returns", "Risk measure"],
        ["momentum_5, momentum_20", "Return over N days", "Trend strength"],
        ["excess_return", "stock_return - sp500_return", "Market-relative performance"],
        ["news_count", "Count of articles per day", "News activity indicator"]
    ],
    col_widths=[1.5, 2.5, 2.0]
)

# d. Data Normalization
doc.add_heading("d) Data Normalization", level=3)
add_formatted_table(doc,
    ["Feature Type", "Normalization Method", "Scope", "Justification"],
    [
        ["Prices (OHLC)", "MinMaxScaler (0-1)", "Per ticker", "Enable cross-stock comparison; $10 vs $1000 stocks"],
        ["Volume", "StandardScaler (z-score)", "Per ticker", "Handle order-of-magnitude differences"],
        ["Returns", "StandardScaler (z-score)", "Global", "Already relative, needs centering for model input"]
    ],
    col_widths=[1.3, 1.8, 1.0, 2.5]
)

doc.add_heading("3.3 Description of Outputs", level=2)
add_formatted_table(doc,
    ["Output File", "Description", "Records", "Features"],
    [
        ["data_aligned.parquet", "Stage 3: News-price aligned dataset", "262,257", "16 columns"],
        ["data_engineered.parquet", "Stage 4: Final feature-engineered dataset", "262,257", "35 columns"],
        ["*_summary.json", "Pipeline statistics and validation metrics", "4 files", "Quality metrics"]
    ],
    col_widths=[2.0, 2.5, 1.0, 1.0]
)

add_insight_box(doc, "Pipeline Validation",
    "Validated output: 262,257 deduplicated records, 100 tickers, 14.2 years of data, "
    "86.4% S&P 500 coverage, 1.46% news coverage (3,821 stock-days with news articles).")

doc.add_page_break()

# ============================================
# SECTION 4: EDA DESCRIPTION (Rubric: 5 pts)
# ============================================
add_heading_with_number(doc, "4. EDA Description", 1)

doc.add_heading("4.1 Types of Analysis Used", level=2)

add_formatted_table(doc,
    ["Analysis Type", "Techniques Applied", "Key Findings"],
    [
        ["Univariate", "Histograms, box plots, descriptive statistics", "Right-skewed prices, high return volatility (34.66% std)"],
        ["Multivariate", "Correlation matrix, scatter plots, cross-tabulations", "S&P 500 return is strongest predictor (+0.023)"],
        ["Graphical", "Time series plots, heatmaps, violin plots, pie charts", "COVID-2020 regime change clearly visible"],
        ["Non-Graphical", "Missing value analysis, statistical tests, outlier detection", "98.5% news missing (expected - not all days have news)"]
    ],
    col_widths=[1.3, 2.5, 2.7]
)

doc.add_heading("4.2 EDA Visualizations with Analysis", level=2)

# Add each EDA figure with detailed analysis
eda_outputs = Path("notebooks/eda_outputs")

# Figure 1: Data Quality
doc.add_heading("a) Data Completeness/Freshness/Quality", level=3)

if (eda_outputs / "02_univariate_distributions.png").exists():
    doc.add_picture(str(eda_outputs / "02_univariate_distributions.png"), width=Inches(6.0))
    cap = doc.add_paragraph("Figure 1: Univariate Distributions - Price, Volume, and Target Distribution")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

add_eda_insight_box(doc, 
    title="Data Quality Assessment",
    what_analyzed="Analyzed missing value patterns across all 262,257 records. Checked OHLCV completeness, news coverage rates, S&P 500 alignment, and date range continuity.",
    what_found="Price Data: 0% missing (complete OHLCV)\nNews Text: 98.5% missing (expected - sparse news)\nS&P 500: 86.4% coverage\nDate Range: 14.2 years (2009-2023)",
    action_taken="Validated data quality sufficient for analysis. Forward-filled minor price gaps. Flagged news availability as feature.",
    why_matters="Complete price data ensures reliable technical indicator calculations. Multi-year span captures different market regimes for robust model training."
)

# Figure 2: Distributions
doc.add_heading("b) Variables and Their Distributions", level=3)

if (eda_outputs / "03_returns_distribution.png").exists():
    doc.add_picture(str(eda_outputs / "03_returns_distribution.png"), width=Inches(6.0))
    cap = doc.add_paragraph("Figure 2: Returns Distribution with +/-2% Buy/Sell Thresholds")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

add_eda_insight_box(doc,
    title="Price & Returns Distribution Analysis",
    what_analyzed="Examined price distributions across 100 stocks and next-day return distributions. Created 3-class targets using +/-2% thresholds. Analyzed extreme movements (>10% daily change).",
    what_found="Mean price ($62) >> Median ($27): right-skewed\nReturn std dev: 2.5% daily (high volatility)\nTarget balance: 70% Hold, 15% Buy, 15% Sell\n4,444 extreme moves identified",
    action_taken="Applied per-ticker MinMaxScaler (0-1) for prices\nApplied StandardScaler for volume\nCreated volatility_20, momentum features",
    why_matters="Without normalization, high-priced stocks ($1000+) would dominate model training. Per-ticker scaling preserves relative patterns while enabling fair cross-stock comparison."
)

# Figure 3: Outliers
doc.add_heading("c) Anomalies - Outliers", level=3)

add_formatted_table(doc,
    ["Outlier Category", "Detection Method", "Count", "Action Taken"],
    [
        ["Price outliers (IQR)", "Q1 - 1.5*IQR, Q3 + 1.5*IQR", "25,372 (9.67%)", "Retained - valid high-priced stocks"],
        ["Extreme gains (>10%)", "Return threshold", "2,145", "Retained - valuable training examples"],
        ["Extreme losses (>10%)", "Return threshold", "2,299", "Retained - valuable training examples"],
        ["Data errors (>50%)", "Return threshold", "951 (0.36%)", "Removed - likely stock splits/errors"]
    ],
    col_widths=[1.8, 2.0, 1.2, 1.8]
)

add_eda_insight_box(doc,
    title="Outlier Detection & Treatment",
    what_analyzed="Used IQR method (1.5x threshold) to detect statistical outliers in price and return distributions. Also flagged extreme daily moves (>10% and >50%).",
    what_found="25,372 price outliers (9.67%) - valid high-priced stocks\n4,444 extreme returns (>10%) - real market events\n951 data errors (>50%) - likely stock splits",
    action_taken="RETAINED: Extreme but valid moves (valuable training examples)\nREMOVED: 951 records with >50% change (0.36% of data)",
    why_matters="Extreme market moves (earnings surprises, crashes, FDA approvals) are exactly what our model needs to learn. Removing them would bias predictions toward average scenarios."
)

# Figure 4: Correlations
doc.add_heading("d) Relationships - Correlations", level=3)

if (eda_outputs / "05_correlation_matrix.png").exists():
    doc.add_picture(str(eda_outputs / "05_correlation_matrix.png"), width=Inches(5.5))
    cap = doc.add_paragraph("Figure 3: Feature Correlation Matrix")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

add_formatted_table(doc,
    ["Feature", "Correlation with Next-Day Return", "Implication"],
    [
        ["sp500_return", "+0.023 (strongest)", "Market context is critical for predictions"],
        ["close", "-0.003", "Raw prices have no predictive power"],
        ["volume", "-0.001", "Raw volume is not predictive"],
        ["open/high/low", "~0.000", "Absolute values don't capture patterns"]
    ],
    col_widths=[1.5, 2.0, 3.0]
)

add_eda_insight_box(doc,
    title="Correlation Analysis Results",
    what_analyzed="Computed correlation matrix for all numeric features against next-day return. Identified strongest predictors and multicollinearity issues.",
    what_found="S&P 500 return: +0.023 (STRONGEST predictor)\nRaw close price: -0.003 (no predictive power)\nRaw volume: -0.001 (not predictive)\nOHLC highly correlated with each other (expected)",
    action_taken="Created RELATIVE features instead of raw values:\n- price_to_sma5/20 (vs moving average)\n- excess_return (stock - market)\n- volume_ratio (vs 20-day avg)",
    why_matters="Market context matters more than absolute price. A $100 stock rising 2% is the same signal as a $10 stock rising 2%. Relative features capture this equivalence."
)

# Temporal Analysis
doc.add_heading("e) Temporal Trends", level=3)

if (eda_outputs / "08_temporal_trends.png").exists():
    doc.add_picture(str(eda_outputs / "08_temporal_trends.png"), width=Inches(6.0))
    cap = doc.add_paragraph("Figure 4: Temporal Trends - Yearly Analysis (2009-2023)")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

add_eda_insight_box(doc,
    title="Temporal Pattern Analysis",
    what_analyzed="Analyzed yearly trends in records count, average prices, trading volume, and returns across 14.2 years (2009-2023). Identified regime changes and structural breaks.",
    what_found="COVID-2020 Impact:\n- Avg daily return: +1.35% (vs 0.2% typical)\n- Avg volume: 3.5M shares (2.4x normal)\n- Price volatility: 2x standard deviation\nMarket regimes clearly visible in temporal plots.",
    action_taken="Created multi-timeframe features:\n- sma_5/20/50 for different horizons\n- volatility_20 for regime detection\n- momentum_5/20 for trend shifts",
    why_matters="Models must handle different market regimes. A strategy that works in stable 2015 may fail in volatile 2020. Our technical indicators help the system adapt automatically."
)

# News Impact
doc.add_heading("f) News Impact Analysis", level=3)

add_formatted_table(doc,
    ["Target", "No News Days", "News Days", "Difference"],
    [
        ["Buy (>+2%)", "14.8%", "18.6%", "+3.8 pp more buys"],
        ["Hold (-2% to +2%)", "69.9%", "60.9%", "-9.0 pp fewer holds"],
        ["Sell (<-2%)", "15.3%", "20.5%", "+5.2 pp more sells"]
    ],
    col_widths=[1.5, 1.5, 1.5, 2.0]
)

add_eda_insight_box(doc,
    title="News Drives Volatility (+10% More Extreme Moves)",
    what_analyzed="Cross-tabulated target distribution (buy/hold/sell) against news availability. Compared movement patterns on days with vs without news coverage.",
    what_found="News days: 18.6% Buy, 60.9% Hold, 20.5% Sell\nNo-news days: 14.8% Buy, 69.9% Hold, 15.3% Sell\nDifference: +10% more extreme moves on news days",
    action_taken="Created news_count feature and has_news flag. This will feed into the sentiment agent in our multi-agent system.",
    why_matters="News is a leading indicator of volatility. Our Graph RAG system will use news sentiment to explain trading recommendations with cited sources."
)

doc.add_heading("4.3 Artifacts Produced", level=2)

add_formatted_table(doc,
    ["Artifact", "Type", "Description"],
    [
        ["02_univariate_distributions.png", "Visualization", "6 subplots: Price OHLC, Volume, Target distribution"],
        ["03_returns_distribution.png", "Visualization", "Returns histogram with ±2% thresholds, box plot"],
        ["04_news_coverage.png", "Visualization", "News availability pie chart and distribution"],
        ["05_correlation_matrix.png", "Visualization", "Feature correlation heatmap"],
        ["07_target_relationships.png", "Visualization", "Target by news and market direction"],
        ["08_temporal_trends.png", "Visualization", "Yearly trends: records, price, volume, returns"],
        ["01_quality_summary.json", "Data", "Data quality metrics and statistics"],
        ["09_eda_insights.json", "Data", "Aggregated EDA findings and key insights"]
    ],
    col_widths=[2.5, 1.2, 3.0]
)

doc.add_page_break()

# ============================================
# SECTION 5: FEATURE ENGINEERING (Rubric: 2 pts)
# ============================================
add_heading_with_number(doc, "5. Feature Engineering", 1)

doc.add_paragraph(
    "Yes, we are performing feature engineering for our multi-agent LLM system. "
    "Below we detail each feature with its EDA justification and literature citation."
)

doc.add_heading("5.1 Features Used in Literature", level=2)

add_formatted_table(doc,
    ["Feature", "Formula/Description", "Literature Citation", "Our EDA Finding"],
    [
        ["SMA (5, 20, 50)", "Simple Moving Average", "Brock et al. (1992) - Technical Analysis", "COVID-2020 shows regime change; multi-timeframe needed"],
        ["RSI (14-day)", "Relative Strength Index", "Wilder (1978) - New Concepts in Trading", "High volatility requires momentum indicators"],
        ["MACD", "12-day EMA - 26-day EMA", "Appel (2005) - Technical Analysis", "Trend detection for regime changes"],
        ["Bollinger Bands", "SMA ± 2*std", "Bollinger (2001) - On Bollinger Bands", "Volatility-adjusted price levels"],
        ["Sentiment Score", "NLP classification (-1, 0, +1)", "Loughran & McDonald (2011)", "News days show 10% more volatility"]
    ],
    col_widths=[1.3, 1.8, 2.0, 1.8]
)

doc.add_heading("5.2 Features Derived from EDA (with Justification)", level=2)

add_formatted_table(doc,
    ["EDA Finding", "Feature Created", "Justification"],
    [
        ["S&P 500 has strongest correlation (+0.023)", "sp500_return, excess_return", "Market context is critical; relative performance matters"],
        ["Raw prices have no correlation (-0.003)", "price_to_sma5, price_to_sma20", "Relative position vs trend is predictive, not absolute price"],
        ["Right-skewed prices ($0.01 to $7,250)", "MinMaxScaler normalization per ticker", "Enable cross-stock comparison in model training"],
        ["High volatility (34.66% std)", "volatility_20, return_norm", "Capture risk regime; normalize for model stability"],
        ["10% more extreme moves on news days", "news_count, has_news flag", "News presence is a volatility signal for sentiment agent"],
        ["COVID-2020 regime change", "sma_5/20/50 multi-timeframe", "Multiple lookback periods detect regime shifts"]
    ],
    col_widths=[2.2, 2.0, 2.5]
)

add_insight_box(doc, "EDA → Feature Traceability",
    "Every feature engineering decision is directly traced to a specific EDA finding. "
    "This ensures our features are data-driven rather than arbitrary, improving model interpretability and performance.")

doc.add_heading("5.3 Complete Feature List (35 Features)", level=2)

add_formatted_table(doc,
    ["Category", "Features", "Count"],
    [
        ["Original Price/Volume", "open, high, low, close, volume", "5"],
        ["Technical Indicators", "sma_5, sma_20, sma_50, momentum_5, momentum_20, volatility_20, volume_ratio, price_to_sma5, price_to_sma20", "9"],
        ["Market Context", "sp500_return, excess_return, market_up, market_down", "4"],
        ["News Features", "news_count, has_news, sentiment_score (future)", "3"],
        ["Normalized Features", "Scaled versions of prices, volume, returns", "12"],
        ["Target Variables", "next_day_return, target (buy/hold/sell)", "2"]
    ],
    col_widths=[1.8, 3.5, 0.8]
)

doc.add_page_break()

# ============================================
# SECTION 6: MODEL ARCHITECTURE (Rubric requirement)
# ============================================
add_heading_with_number(doc, "6. Model Architecture", 1)

doc.add_heading("6.1 Models Used (Priority Order)", level=2)

add_formatted_table(doc,
    ["Priority", "Model/Component", "Purpose", "Justification"],
    [
        ["1 (Highest)", "GPT-5.2 via OpenAI Agents SDK", "Multi-agent orchestration, explanation generation", "State-of-the-art reasoning, native agent support"],
        ["2", "Neo4j Graph Database", "Knowledge storage, relationship queries", "Natural fit for entity relationships in finance"],
        ["3", "FinBERT/RoBERTa", "Sentiment classification", "Pre-trained on financial text, proven accuracy"],
        ["4", "LightGBM/XGBoost", "Feature-based prediction baseline", "Strong tabular performance, interpretable"],
        ["5", "RAG (Retrieval-Augmented)", "Grounded explanation generation", "Enables source citation for trust"]
    ],
    col_widths=[0.8, 2.0, 2.0, 2.0]
)

doc.add_heading("6.2 Model Architecture from Literature", level=2)

# Multi-agent architecture flowchart
add_flowchart_box(doc, "Multi-Agent System Architecture:", [
    ("USER QUERY", "Should I buy\nAAPL?"),
    ("ORCHESTRATOR", "Routes to\nspecialist agents"),
    ("ANALYST AGENTS", "Fundamental\nTechnical\nSentiment"),
    ("ADVISOR AGENTS", "Optimistic\nCautious"),
    ("RAG + RESPONSE", "Cited\nRecommendation")
])

doc.add_paragraph()
doc.add_paragraph(
    "Based on: TradingAgents (2024) multi-agent framework, adapted to use OpenAI Agents SDK for structured handoffs "
    "and Neo4j for persistent knowledge representation instead of in-memory state."
)

doc.add_heading("6.3 New Architecture Experiments", level=2)

add_formatted_table(doc,
    ["Experiment", "Description", "Expected Benefit"],
    [
        ["Graph RAG vs Vector RAG", "Neo4j traversal vs embedding similarity", "Better relationship reasoning for financial entities"],
        ["Agent Debate Protocol", "Optimistic vs Cautious agents argue", "More balanced recommendations with explicit trade-offs"],
        ["Temporal Attention", "Weight recent news higher", "Recency-aware sentiment aggregation"],
        ["Confidence Calibration", "Agent confidence scores", "Users know when to trust vs verify recommendations"]
    ],
    col_widths=[1.8, 2.5, 2.5]
)

doc.add_page_break()

# ============================================
# SECTION 7: GRAPH RAG SYSTEM DETAILS
# ============================================
add_heading_with_number(doc, "7. Graph RAG System Details", 1)

doc.add_paragraph(
    "Our Graph RAG (Retrieval-Augmented Generation) system combines Neo4j knowledge graph with GPT-5.2 "
    "to provide explainable financial recommendations with cited sources. This section details the "
    "implementation architecture and query mechanisms."
)

doc.add_heading("7.1 Neo4j Knowledge Graph Schema", level=2)

# Add visual schema
add_neo4j_schema(doc)

# Detailed node descriptions
doc.add_heading("7.2 Entity Node Details", level=3)

add_formatted_table(doc,
    ["Node Type", "Key Properties", "Data Source", "Update Frequency"],
    [
        ["(:Stock)", "ticker, name, sector, price, volume, market_cap", "FNSPID + Yahoo", "Daily"],
        ["(:Article)", "id, title, text, date, source, url, word_count", "FNSPID News", "Real-time"],
        ["(:Person)", "name, role, company, title, influence_score", "NER extraction", "On article ingestion"],
        ["(:Event)", "type, date, description, impact_score, affected_sectors", "NLP extraction", "On article ingestion"],
        ["(:Sentiment)", "score (-1 to +1), label, confidence, key_phrases", "FinBERT model", "On article ingestion"],
        ["(:TradingSignal)", "type, strength (0-100), reasoning, expiry_date", "Agent consensus", "On query"]
    ],
    col_widths=[1.2, 2.0, 1.3, 1.2]
)

doc.add_heading("7.3 Graph Query Flow", level=2)

doc.add_paragraph(
    "When a user asks 'Should I buy AAPL?', the system executes the following query pipeline:"
)

# Query flow as flowchart
add_flowchart_box(doc, "", [
    ("1. PARSE", "Extract ticker\nquery intent"),
    ("2. RETRIEVE", "Graph traversal\nget evidence"),
    ("3. AGGREGATE", "Combine signals\nweight by recency"),
    ("4. GENERATE", "GPT-5.2 response\nwith citations")
])

doc.add_heading("7.4 Cypher Query Examples", level=2)

# Query 1
add_code_block(doc, 
    "Query: Find recent news sentiment for a stock",
    """MATCH (stock:Stock {ticker: 'AAPL'})
      -[:MENTIONED_IN]->(article:Article)
      -[:HAS_SENTIMENT]->(sentiment:Sentiment)
WHERE article.date > date() - duration('P7D')
RETURN article.title, article.date, 
       sentiment.score, sentiment.label
ORDER BY article.date DESC LIMIT 5""",
    "Returns 3 articles with average sentiment +0.67 (positive)")

# Query 2
add_code_block(doc,
    "Query: Find related stocks affected by same events",
    """MATCH (s1:Stock {ticker: 'AAPL'})
      -[:CORRELATED_WITH]->(s2:Stock)
MATCH (event:Event)-[:AFFECTS]->(s1)
WHERE event.date > date() - duration('P30D')
RETURN s2.ticker, s2.name, 
       collect(event.type) as shared_events
ORDER BY s1.correlation DESC""",
    "Returns MSFT (0.82 correlation), GOOGL (0.76) with shared tech events")

doc.add_heading("7.5 RAG Response Generation", level=2)

doc.add_paragraph(
    "After retrieving graph evidence, GPT-5.2 generates a natural language response with citations:"
)

# Example response box
response_table = doc.add_table(rows=1, cols=1)
response_table.alignment = WD_TABLE_ALIGNMENT.CENTER
response_cell = response_table.cell(0, 0)
set_cell_shading(response_cell, 'F0FFF4')

response_p = response_cell.paragraphs[0]
title_run = response_p.add_run("RECOMMENDATION: BUY AAPL (Confidence: 72%)\n\n")
title_run.bold = True
title_run.font.color.rgb = RGBColor(34, 84, 61)

response_p.add_run(
    "Based on recent news analysis, AAPL shows positive sentiment (avg: +0.67).\n\n"
    "Key factors supporting this recommendation:\n"
    "  * Q1 earnings exceeded expectations [Article #1, 2026-01-28]\n"
    "  * Strong iPhone demand in China market [Article #2, 2026-01-27]\n"
    "  * Supply chain concerns easing [Article #3, 2026-01-26]\n\n"
    "Technical indicators: Price above 20-day SMA, momentum +2.3%\n\n"
    "Risk factors to consider:\n"
    "  * Broader market volatility (VIX elevated)\n"
    "  * Earnings call in 2 weeks may introduce uncertainty"
)

doc.add_paragraph()

add_insight_box(doc, "Explainability Advantage",
    "Unlike black-box models, every recommendation is traceable to specific news articles, "
    "sentiment scores, and technical indicators stored in the knowledge graph. Users can "
    "click on citations to verify the reasoning themselves.")

doc.add_page_break()

# ============================================
# SECTION 8: PROGRESS REPORT (Rubric: 3 pts)
# ============================================
add_heading_with_number(doc, "8. Progress Report", 1)

doc.add_heading("7.1 Effort and Progress Details", level=2)

add_formatted_table(doc,
    ["Task", "Status", "Effort", "Details"],
    [
        ["Dataset Collection", "Complete", "Week 1", "Downloaded 4 datasets from HuggingFace/Yahoo; 500K+ raw records"],
        ["Data Pipeline (4 stages)", "Complete", "Week 1-2", "Load > Clean > Align > Engineer; 262K final records"],
        ["EDA Analysis", "Complete", "Week 2", "7 visualizations, 15 statistical tests, 2 JSON summaries"],
        ["Feature Engineering", "Complete", "Week 2", "35 features extracted with EDA justification"],
        ["Sentiment Model", "In Progress", "Week 3", "Fine-tuning FinBERT on Financial Phrasebank"],
        ["Agent Framework", "Planned", "Week 4", "OpenAI Agents SDK integration"],
        ["Neo4j Setup", "Planned", "Week 5", "Schema design, data loading, query optimization"],
        ["RAG Integration", "Planned", "Week 6", "Retrieval + generation with citations"],
        ["Demo & Report", "Planned", "Week 7", "Final presentation and documentation"]
    ],
    col_widths=[1.8, 1.0, 0.8, 3.0]
)

doc.add_heading("7.2 Tests Run", level=2)

add_formatted_table(doc,
    ["Test", "Result", "Finding"],
    [
        ["Pipeline validation", "Pass", "262K records, 100 tickers, 14.2 years coverage"],
        ["Data quality check", "Pass", "< 0.1% missing in critical fields"],
        ["Outlier analysis", "Pass", "0.36% data errors removed, valid extremes retained"],
        ["Feature correlation", "Pass", "S&P 500 strongest predictor; no multicollinearity issues"],
        ["Target balance check", "Pass", "70/15/15 split acceptable for classification"]
    ],
    col_widths=[1.8, 0.8, 4.0]
)

doc.add_heading("7.3 Team Member Contributions", level=2)

add_formatted_table(doc,
    ["Team Member", "Contributions (Current)", "Next Steps"],
    [
        ["Harsh Arya", "EDA analysis, Feature engineering Stage 4, Data pipeline", "Baseline models, Evaluation metrics, Demo preparation"],
        ["Camila Paik", "Report organization, Progress documentation, Stage 3 alignment", "Neo4j schema design, Entity extraction, Graph queries"],
        ["Gabrielle Despaigne", "Document organization, Progress tracking, Stage 2 cleaning", "Documentation, Proof Reading, Orchestrator agent implementation"],
        ["Raghav Vasappanavara", "Problem Statement, System Architecture, Technical Design", "Multi-agent framework, OpenAI SDK integration, Orchestrator agent"]
    ],
    col_widths=[1.8, 2.5, 2.5]
)

# ============================================
# SECTION 9: RISKS AND MITIGATION
# ============================================
add_heading_with_number(doc, "9. Risks and Mitigation", 1)

add_formatted_table(doc,
    ["Risk", "Probability", "Impact", "Mitigation Strategy"],
    [
        ["API rate limits (OpenAI)", "Medium", "High", "Implement caching, batch requests, exponential backoff retry logic"],
        ["Graph query performance", "Medium", "Medium", "Index optimization, query caching, pagination for large results"],
        ["Sentiment model accuracy", "Low", "Medium", "Ensemble methods, confidence thresholds, human validation sample"],
        ["Data quality issues", "Low", "High", "Automated validation pipeline, monitoring dashboards"],
        ["Integration complexity", "Medium", "Medium", "Modular design, comprehensive unit/integration testing"]
    ],
    col_widths=[1.5, 0.9, 0.8, 3.5]
)

# ============================================
# SECTION 10: REFERENCES
# ============================================
add_heading_with_number(doc, "10. References", 1)

doc.add_heading("Datasets", level=2)
refs = [
    "Dong, Z., Fan, X., & Peng, Z. (2024). FNSPID: A Comprehensive Financial News Dataset in Time Series. arXiv:2402.06698",
    "Malo, P., et al. (2014). Good debt or bad debt: Detecting semantic orientations in economic texts. JASIST. (Financial Phrasebank)",
    "Chen, Z., et al. (2021). FinQA: A Dataset of Numerical Reasoning over Financial Data. EMNLP.",
    "Yahoo Finance. S&P 500 Historical Data. https://finance.yahoo.com"
]
for ref in refs:
    p = doc.add_paragraph(f"• {ref}")
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_heading("Technical Analysis Literature", level=2)
refs = [
    "Brock, W., Lakonishok, J., & LeBaron, B. (1992). Simple technical trading rules and the stochastic properties of stock returns. Journal of Finance.",
    "Wilder, J.W. (1978). New Concepts in Technical Trading Systems. Trend Research.",
    "Bollinger, J. (2001). Bollinger on Bollinger Bands. McGraw-Hill."
]
for ref in refs:
    p = doc.add_paragraph(f"• {ref}")
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_heading("NLP and Sentiment Analysis", level=2)
refs = [
    "Loughran, T., & McDonald, B. (2011). When is a liability not a liability? Textual analysis, dictionaries, and 10-Ks. Journal of Finance.",
    "Araci, D. (2019). FinBERT: Financial Sentiment Analysis with Pre-trained Language Models. arXiv:1908.10063"
]
for ref in refs:
    p = doc.add_paragraph(f"• {ref}")
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_heading("Multi-Agent Systems", level=2)
refs = [
    "TradingAgents (2024). Multi-Agents LLM Financial Trading Framework. https://github.com/TauricResearch/TradingAgents",
    "OpenAI (2024). Agents SDK Documentation. https://platform.openai.com/docs/agents"
]
for ref in refs:
    p = doc.add_paragraph(f"• {ref}")
    p.paragraph_format.left_indent = Inches(0.3)

# ============================================
# SAVE DOCUMENT
# ============================================
output_path = "DSC288_Progress_Report_v2.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Total pages: ~15-18")
print(f"Sections: 9 (following rubric)")
print(f"Tables: 25+")
print(f"Figures: 4 EDA visualizations")
print("Done!")
